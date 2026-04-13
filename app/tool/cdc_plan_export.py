import json
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from docx import Document
from docx.shared import Pt

from app.config import config
from app.schema import (
    CDCMeasureLevel,
    CDCPlanDocument,
    CDCPlanMeta,
    CDCPlanSection,
)
from app.tool.base import BaseTool, ToolResult


class CDCPlanExportTool(BaseTool):
    name: str = "cdc_plan_export"
    description: str = "Export a CDC emergency plan to a Word .docx file."
    parameters: dict = {
        "type": "object",
        "properties": {
            "plan": {
                "description": "Plan object or JSON string that matches CDCPlanDocument schema",
                "anyOf": [{"type": "object"}, {"type": "string"}],
            },
            "output_path": {
                "type": "string",
                "description": "Output path. If relative, it is created under workspace.",
            },
        },
        "required": ["plan"],
    }

    @staticmethod
    def _normalize_output_path(output_path: Optional[str]) -> Path:
        if output_path:
            p = Path(output_path)
            if p.is_absolute():
                return p
            return config.workspace_root / p
        ts = int(time.time())
        return config.workspace_root / f"cdc_plan_{ts}.docx"

    @staticmethod
    def _parse_plan(plan: Union[str, Dict[str, Any]]) -> CDCPlanDocument:
        if isinstance(plan, str):
            raw = plan.strip()
            if not raw:
                raise ValueError("plan is empty")
            data = json.loads(raw)
            if not isinstance(data, dict):
                raise ValueError("plan JSON must be an object")
            normalized = CDCPlanExportTool._normalize_plan_dict(data)
            return CDCPlanDocument(**normalized)
        if isinstance(plan, dict):
            normalized = CDCPlanExportTool._normalize_plan_dict(plan)
            return CDCPlanDocument(**normalized)
        raise ValueError("plan must be a dict or JSON string")

    @staticmethod
    def _first_nonempty(*values: Any, default: Any = None) -> Any:
        for v in values:
            if v is None:
                continue
            if isinstance(v, str) and not v.strip():
                continue
            return v
        return default

    @staticmethod
    def _as_int(value: Any, default: int) -> int:
        if value is None:
            return default
        try:
            if isinstance(value, bool):
                return default
            if isinstance(value, (int, float)):
                return int(value)
            s = str(value).strip()
            if not s:
                return default
            return int(float(s))
        except Exception:
            return default

    @staticmethod
    def _as_float(value: Any, default: float) -> float:
        if value is None:
            return default
        try:
            if isinstance(value, bool):
                return default
            if isinstance(value, (int, float)):
                return float(value)
            s = str(value).strip()
            if not s:
                return default
            return float(s)
        except Exception:
            return default

    @staticmethod
    def _normalize_event_type(raw: Any) -> str:
        if raw is None:
            return "other"
        v = str(raw).strip()
        if not v:
            return "other"
        mapping = {
            "学校流感": "influenza_school",
            "流感": "influenza_school",
            "社区新冠": "covid_community",
            "新冠": "covid_community",
            "诺如": "norovirus_cluster",
            "诺如病毒": "norovirus_cluster",
            "诺如病毒聚集": "norovirus_cluster",
        }
        if v in mapping:
            return mapping[v]
        return v

    @staticmethod
    def _normalize_risk_level(raw: Any) -> str:
        if raw is None:
            return "low"
        v = str(raw).strip().lower()
        if not v:
            return "low"
        mapping = {
            "低": "low",
            "低风险": "low",
            "中": "medium",
            "中风险": "medium",
            "高": "high",
            "高风险": "high",
            "极高": "extreme",
            "极高风险": "extreme",
        }
        if v in mapping:
            return mapping[v]
        return v

    @staticmethod
    def _normalize_sections(raw_sections: Any) -> List[Dict[str, Any]]:
        if not isinstance(raw_sections, list):
            return []

        def normalize_section(sec: Any) -> Dict[str, Any]:
            if not isinstance(sec, dict):
                return {"title": str(sec), "paragraphs": [], "subsections": []}

            title = CDCPlanExportTool._first_nonempty(
                sec.get("title"),
                sec.get("section_title"),
                sec.get("sectionTitle"),
                sec.get("subsection_title"),
                sec.get("subsectionTitle"),
                default="未命名章节",
            )

            paragraphs_val = CDCPlanExportTool._first_nonempty(
                sec.get("paragraphs"),
                sec.get("paras"),
                sec.get("content"),
                sec.get("text"),
                default=[],
            )
            paragraphs: List[str] = []
            if isinstance(paragraphs_val, list):
                paragraphs = [str(p) for p in paragraphs_val if str(p).strip()]
            elif isinstance(paragraphs_val, str):
                p = paragraphs_val.strip()
                if p:
                    paragraphs = [p]

            raw_subsections = CDCPlanExportTool._first_nonempty(
                sec.get("subsections"),
                sec.get("sub_sections"),
                sec.get("children"),
                default=[],
            )
            subsections: List[Dict[str, Any]] = []
            if isinstance(raw_subsections, list):
                subsections = [normalize_section(s) for s in raw_subsections]

            return {"title": str(title), "paragraphs": paragraphs, "subsections": subsections}

        return [normalize_section(s) for s in raw_sections]

    @staticmethod
    def _normalize_citations(raw: Any) -> List[Dict[str, Any]]:
        if not isinstance(raw, list):
            return []
        out: List[Dict[str, Any]] = []
        for c in raw:
            if not isinstance(c, dict):
                continue
            out.append(
                {
                    "source_file": str(
                        CDCPlanExportTool._first_nonempty(
                            c.get("source_file"),
                            c.get("source"),
                            c.get("file"),
                            c.get("filename"),
                            default="unknown",
                        )
                    ),
                    "chunk_id": CDCPlanExportTool._as_int(
                        CDCPlanExportTool._first_nonempty(
                            c.get("chunk_id"),
                            c.get("chunkId"),
                            c.get("id"),
                            default=0,
                        ),
                        0,
                    ),
                    "score": CDCPlanExportTool._as_float(
                        CDCPlanExportTool._first_nonempty(
                            c.get("score"),
                            c.get("similarity"),
                            c.get("relevance"),
                            default=0.0,
                        ),
                        0.0,
                    ),
                    "excerpt": str(
                        CDCPlanExportTool._first_nonempty(
                            c.get("excerpt"),
                            c.get("text"),
                            c.get("content"),
                            default="",
                        )
                    ),
                }
            )
        return out

    @staticmethod
    def _normalize_measures(raw_measures: Any) -> List[Dict[str, Any]]:
        if not isinstance(raw_measures, list):
            return []
        out: List[Dict[str, Any]] = []
        for m in raw_measures:
            if not isinstance(m, dict):
                continue
            level_raw = CDCPlanExportTool._first_nonempty(
                m.get("level"),
                m.get("measure_level"),
                m.get("type"),
                default="core",
            )
            level = str(level_raw).strip().lower()
            if level in {"核心", "强合规"}:
                level = "core"
            if level in {"补充", "可选", "适配"}:
                level = "supplementary"

            out.append(
                {
                    "title": str(
                        CDCPlanExportTool._first_nonempty(
                            m.get("title"),
                            m.get("measure_title"),
                            m.get("name"),
                            default="未命名措施",
                        )
                    ),
                    "content": str(
                        CDCPlanExportTool._first_nonempty(
                            m.get("content"),
                            m.get("measure_content"),
                            m.get("description"),
                            m.get("text"),
                            default="",
                        )
                    ),
                    "level": level,
                    "citations": CDCPlanExportTool._normalize_citations(
                        CDCPlanExportTool._first_nonempty(
                            m.get("citations"),
                            m.get("references"),
                            m.get("refs"),
                            default=[],
                        )
                    ),
                }
            )
        return out

    @staticmethod
    def _normalize_resources(raw: Any) -> Dict[str, Any]:
        if isinstance(raw, dict):
            items = raw.get("items")
            if isinstance(items, list):
                return {"items": CDCPlanExportTool._normalize_resource_items(items)}
            if isinstance(raw.get("materials"), list):
                return {"items": CDCPlanExportTool._normalize_resource_items(raw.get("materials"))}
        if isinstance(raw, list):
            return {"items": CDCPlanExportTool._normalize_resource_items(raw)}
        return {"items": []}

    @staticmethod
    def _normalize_resource_items(raw_items: Any) -> List[Dict[str, Any]]:
        if not isinstance(raw_items, list):
            return []
        out: List[Dict[str, Any]] = []
        for it in raw_items:
            if not isinstance(it, dict):
                continue
            out.append(
                {
                    "name": str(
                        CDCPlanExportTool._first_nonempty(
                            it.get("name"),
                            it.get("item_name"),
                            it.get("resource_name"),
                            default="未知物资",
                        )
                    ),
                    "unit": str(
                        CDCPlanExportTool._first_nonempty(it.get("unit"), default="unit")
                    ),
                    "quantity": max(
                        0.0,
                        CDCPlanExportTool._as_float(
                            CDCPlanExportTool._first_nonempty(
                                it.get("quantity"),
                                it.get("count"),
                                it.get("amount"),
                                default=0.0,
                            ),
                            0.0,
                        ),
                    ),
                }
            )
        return out

    @staticmethod
    def _normalize_plan_dict(data: Dict[str, Any]) -> Dict[str, Any]:
        if not isinstance(data, dict):
            return {}

        meta_in = data.get("meta") if isinstance(data.get("meta"), dict) else {}
        input_in = data.get("input") if isinstance(data.get("input"), dict) else {}
        risk_in = data.get("risk") if isinstance(data.get("risk"), dict) else {}

        meta = {
            "title": str(
                CDCPlanExportTool._first_nonempty(
                    meta_in.get("title"),
                    meta_in.get("plan_title"),
                    meta_in.get("planTitle"),
                    data.get("plan_title"),
                    data.get("planTitle"),
                    data.get("title"),
                    default="疾控应急预案",
                )
            ),
            "jurisdiction": CDCPlanExportTool._first_nonempty(
                meta_in.get("jurisdiction"),
                meta_in.get("unit"),
                data.get("jurisdiction"),
                data.get("unit"),
                data.get("issuing_unit"),
                default=None,
            ),
            "created_at": CDCPlanExportTool._first_nonempty(
                meta_in.get("created_at"),
                meta_in.get("create_time"),
                data.get("created_at"),
                data.get("create_time"),
                default=None,
            ),
        }

        transmission_in = (
            input_in.get("transmission")
            if isinstance(input_in.get("transmission"), dict)
            else (data.get("transmission") if isinstance(data.get("transmission"), dict) else {})
        )

        event_input = {
            "event_type": CDCPlanExportTool._normalize_event_type(
                CDCPlanExportTool._first_nonempty(
                    input_in.get("event_type"),
                    input_in.get("eventType"),
                    data.get("event_type"),
                    data.get("eventType"),
                    data.get("event"),
                    default="other",
                )
            ),
            "location": str(
                CDCPlanExportTool._first_nonempty(
                    input_in.get("location"),
                    input_in.get("place"),
                    input_in.get("site"),
                    data.get("location"),
                    data.get("place"),
                    data.get("site"),
                    default="未提供",
                )
            ),
            "population": max(
                1,
                CDCPlanExportTool._as_int(
                    CDCPlanExportTool._first_nonempty(
                        input_in.get("population"),
                        input_in.get("region_population"),
                        data.get("population"),
                        data.get("region_population"),
                        default=1,
                    ),
                    1,
                ),
            ),
            "reported_cases": max(
                0,
                CDCPlanExportTool._as_int(
                    CDCPlanExportTool._first_nonempty(
                        input_in.get("reported_cases"),
                        input_in.get("cases"),
                        input_in.get("case_count"),
                        data.get("reported_cases"),
                        data.get("cases"),
                        data.get("case_count"),
                        default=0,
                    ),
                    0,
                ),
            ),
            "report_date": CDCPlanExportTool._first_nonempty(
                input_in.get("report_date"),
                input_in.get("date"),
                data.get("report_date"),
                data.get("date"),
                default=None,
            ),
            "transmission": {
                "r0": transmission_in.get("r0"),
                "incubation_days": transmission_in.get("incubation_days"),
                "infectious_days": transmission_in.get("infectious_days"),
            },
        }

        risk = {
            "level": CDCPlanExportTool._normalize_risk_level(
                CDCPlanExportTool._first_nonempty(
                    risk_in.get("level"),
                    risk_in.get("risk_level"),
                    risk_in.get("riskLevel"),
                    data.get("risk_level"),
                    data.get("riskLevel"),
                    default="low",
                )
            ),
            "summary": str(
                CDCPlanExportTool._first_nonempty(
                    risk_in.get("summary"),
                    risk_in.get("analysis"),
                    risk_in.get("risk_summary"),
                    data.get("risk_summary"),
                    data.get("risk_analysis"),
                    default="未提供风险评估结论。",
                )
            ),
            "predicted_cases_7d": CDCPlanExportTool._first_nonempty(
                risk_in.get("predicted_cases_7d"),
                risk_in.get("prediction_7d"),
                data.get("predicted_cases_7d"),
                data.get("prediction_7d"),
                default=None,
            ),
        }

        sections = CDCPlanExportTool._normalize_sections(
            CDCPlanExportTool._first_nonempty(
                data.get("sections"),
                data.get("plan_sections"),
                data.get("outline"),
                default=[],
            )
        )

        measures = CDCPlanExportTool._normalize_measures(
            CDCPlanExportTool._first_nonempty(
                data.get("measures"),
                data.get("control_measures"),
                data.get("actions"),
                default=[],
            )
        )

        resources = CDCPlanExportTool._normalize_resources(
            CDCPlanExportTool._first_nonempty(
                data.get("resources"),
                data.get("stock"),
                data.get("materials"),
                default={"items": []},
            )
        )

        normalized: Dict[str, Any] = {
            "meta": meta,
            "input": event_input,
            "risk": risk,
            "measures": measures,
            "resources": resources,
            "sections": sections,
        }
        return normalized

    @staticmethod
    def _ensure_sections(plan: CDCPlanDocument) -> List[CDCPlanSection]:
        if plan.sections:
            return plan.sections

        measures_core = [m for m in plan.measures if m.level == CDCMeasureLevel.core]
        measures_supp = [
            m for m in plan.measures if m.level == CDCMeasureLevel.supplementary
        ]

        sections: List[CDCPlanSection] = []
        sections.append(
            CDCPlanSection(
                title="一、事件概况",
                paragraphs=[
                    f"事件类型：{plan.input.event_type}",
                    f"发生地点：{plan.input.location}",
                    f"区域人口：{plan.input.population}",
                    f"报告病例数：{plan.input.reported_cases}",
                ],
            )
        )
        sections.append(
            CDCPlanSection(
                title="二、风险评估",
                paragraphs=[
                    f"风险等级：{plan.risk.level}",
                    f"评估结论：{plan.risk.summary}",
                ]
                + (
                    [f"未来 7 天病例预测：{plan.risk.predicted_cases_7d}"]
                    if plan.risk.predicted_cases_7d is not None
                    else []
                ),
            )
        )
        if measures_core or measures_supp:
            sub = []
            if measures_core:
                sub.append(
                    CDCPlanSection(
                        title="（一）核心措施（强合规）",
                        paragraphs=[
                            f"{i+1}. {m.title}：{m.content}"
                            for i, m in enumerate(measures_core)
                        ],
                    )
                )
            if measures_supp:
                sub.append(
                    CDCPlanSection(
                        title="（二）补充措施（区域适配）",
                        paragraphs=[
                            f"{i+1}. {m.title}：{m.content}"
                            for i, m in enumerate(measures_supp)
                        ],
                    )
                )
            sections.append(CDCPlanSection(title="三、防控措施", subsections=sub))
        sections.append(
            CDCPlanSection(
                title="四、资源与物资保障",
                paragraphs=[
                    f"{i+1}. {item.name}：{item.quantity} {item.unit}"
                    for i, item in enumerate(plan.resources.items)
                ]
                if plan.resources.items
                else ["暂无资源库存数据。"],
            )
        )
        sections.append(
            CDCPlanSection(
                title="五、审批与签字",
                paragraphs=["拟稿：", "审核：", "批准：", "（公章占位）"],
            )
        )
        return sections

    @staticmethod
    def _apply_default_style(doc: Document) -> None:
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(11)

    @staticmethod
    def _write_meta(doc: Document, meta: CDCPlanMeta) -> None:
        doc.add_heading(meta.title, level=0)
        if meta.jurisdiction:
            doc.add_paragraph(f"编制单位：{meta.jurisdiction}")
        if meta.created_at:
            doc.add_paragraph(f"生成时间：{meta.created_at}")
        doc.add_paragraph("")

    @staticmethod
    def _write_section(doc: Document, section: CDCPlanSection, level: int) -> None:
        heading_level = max(1, min(4, level))
        doc.add_heading(section.title, level=heading_level)
        for p in section.paragraphs or []:
            doc.add_paragraph(p)
        for sub in section.subsections or []:
            CDCPlanExportTool._write_section(doc, sub, level + 1)

    @staticmethod
    def _write_citations(doc: Document, plan: CDCPlanDocument) -> None:
        citations = []
        for m in plan.measures:
            for c in m.citations:
                citations.append((m.title, c))
        if not citations:
            return
        doc.add_heading("附：规范依据摘录（引用）", level=1)
        for i, (measure_title, c) in enumerate(citations, 1):
            doc.add_paragraph(
                f"{i}. 措施：{measure_title} | 来源：{c.source_file} | chunk_id：{c.chunk_id} | score：{c.score:.4f}"
            )
            doc.add_paragraph(c.excerpt)

    async def execute(self, **kwargs) -> ToolResult:
        try:
            plan = self._parse_plan(kwargs.get("plan"))
        except Exception as e:
            return ToolResult(error=str(e))

        output_path = self._normalize_output_path(kwargs.get("output_path"))
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        try:
            self._apply_default_style(doc)
        except Exception:
            pass

        self._write_meta(doc, plan.meta)
        sections = self._ensure_sections(plan)
        for s in sections:
            self._write_section(doc, s, level=1)
        self._write_citations(doc, plan)

        doc.save(str(output_path))
        return self.success_response(
            {"output_path": str(output_path), "title": plan.meta.title}
        )
